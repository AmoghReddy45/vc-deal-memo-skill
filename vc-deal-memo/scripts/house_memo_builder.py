#!/usr/bin/env python3
"""Build a house-style VC memo DOCX from a JSON outline.

This script exists because prose instructions are not enough for consistent
memo formatting across models. Give the model a JSON outline, then run this
builder to produce the house DOCX.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "Missing dependency: python-docx. Install it with `python3 -m pip install python-docx`."
    ) from exc


BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
MUTED = RGBColor(0x55, 0x55, 0x55)
DARK_BAR = "070B12"
HEADER_FILL = "D9D9D9"
DEAL_LEFT = "D9D9D9"
DEAL_RIGHT = "EFEFEF"
CALLOUT_FILL = "F4F6F9"
BORDER = "C9D0D8"
PAGE_WIDTH_DXA = 9360


def clean_text(text: Any) -> str:
    value = "" if text is None else str(text)
    return value.replace("\u2014", "-").replace("\u2013", "-")


def set_run_font(run, *, name="Arial", size=10.2, color=BLACK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, *, before=0, after=5, line=1.08):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.type = "dxa"
            tc_w.w = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size="4", inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if edge.startswith("inside") and not inside:
            element.set(qn("w:val"), "nil")
            element.set(qn("w:sz"), "0")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "FFFFFF")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)


def _max_numbering_id(numbering, element_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(f"w:{element_name}")):
        raw = element.get(qn(f"w:{attr_name}"))
        if raw and raw.isdigit():
            values.append(int(raw))
    return max(values, default=-1)


def house_bullet_num_id(doc: Document) -> int:
    existing = getattr(doc, "_house_bullet_num_id", None)
    if existing is not None:
        return existing

    numbering = doc.part.numbering_part.element
    abstract_id = _max_numbering_id(numbering, "abstractNum", "abstractNumId") + 1
    num_id = _max_numbering_id(numbering, "num", "numId") + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (
        ("w:start", "1"),
        ("w:numFmt", "bullet"),
        ("w:lvlText", "\u2022"),
        ("w:lvlJc", "left"),
    ):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        lvl.append(node)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "330")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "330")
    ind.set(qn("w:hanging"), "150")
    ppr.append(ind)
    lvl.append(ppr)

    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(fonts)
    lvl.append(rpr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    setattr(doc, "_house_bullet_num_id", num_id)
    return num_id


def apply_house_bullet(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def keep_table_rows_together(table):
    for r_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if r_idx == 0 and tr_pr.find(qn("w:tblHeader")) is None:
            tr_pr.append(OxmlElement("w:tblHeader"))


def configure_document(doc: Document, footer_text: str):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    footer = section.footer.paragraphs[0]
    footer.text = clean_text(footer_text)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(footer, after=0, line=1.0)
    if footer.runs:
        set_run_font(footer.runs[0], size=8.2, color=MUTED)


def add_runs_from_markdown(paragraph, text: str, *, size=10.2, color=BLACK, bold_default=False):
    parts = re.split(r"(\*\*[^*]+\*\*)", clean_text(text))
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = paragraph.add_run(content)
        set_run_font(run, size=size, color=color, bold=bold_default or is_bold)


def add_paragraph(doc: Document, text: str, *, size=10.2, before=0, after=5, bold=False):
    p = doc.add_paragraph()
    set_para_spacing(p, before=before, after=after, line=1.08)
    add_runs_from_markdown(p, text, size=size, bold_default=bold)
    return p


def add_bullet(doc: Document, item: Any, *, size=9.75):
    p = doc.add_paragraph()
    apply_house_bullet(p, house_bullet_num_id(doc))
    set_para_spacing(p, after=3, line=1.05)
    if isinstance(item, dict):
        label = clean_text(item.get("label", "")).rstrip(":")
        text = clean_text(item.get("text", ""))
        if label:
            r = p.add_run(f"{label}: ")
            set_run_font(r, size=size, bold=True)
        add_runs_from_markdown(p, text, size=size)
    else:
        add_runs_from_markdown(p, clean_text(item), size=size)


def add_section_bar(doc: Document, title: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=DARK_BAR, size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, DARK_BAR)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=0, line=1.0)
    r = p.add_run(clean_text(title).upper())
    set_run_font(r, size=10.8, color=WHITE, bold=True)
    doc.add_paragraph()


def add_masthead(doc: Document, data: dict[str, Any]):
    if data.get("confidential", True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_para_spacing(p, after=8, line=1.0)
        r = p.add_run("[CONFIDENTIAL - DO NOT FORWARD]")
        set_run_font(r, size=9.8, bold=True)

    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=2, line=1.0)
    r = p.add_run(clean_text(data.get("title", "Deal Memo")))
    set_run_font(r, size=19, color=BLACK, bold=True)

    subtitle = clean_text(data.get("subtitle", ""))
    if subtitle:
        sub = doc.add_paragraph()
        set_para_spacing(sub, after=8, line=1.05)
        sr = sub.add_run(subtitle)
        set_run_font(sr, size=10.2, color=MUTED)

    for row in data.get("meta", []):
        label = clean_text(row.get("label", ""))
        value = clean_text(row.get("value", ""))
        if not label and not value:
            continue
        p = doc.add_paragraph()
        set_para_spacing(p, after=1, line=1.0)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.2, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.2)

    rule = doc.add_paragraph()
    set_para_spacing(rule, before=4, after=9, line=1.0)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def infer_widths(headers: list[str], rows: list[list[Any]], explicit: list[int] | None = None) -> list[int]:
    if explicit:
        return explicit
    n = len(headers)
    joined = " ".join(headers).lower()
    if n == 2:
        return [2100, PAGE_WIDTH_DXA - 2100]
    if n == 3:
        return [2300, 2300, PAGE_WIDTH_DXA - 4600]
    if n == 4 and ("read" in joined or "interpretation" in joined):
        return [2050, 2400, 2400, 2510]
    if n == 5:
        return [2100, 1700, 1700, 1700, 2160]
    return [PAGE_WIDTH_DXA // n] * n


def add_table(doc: Document, spec: dict[str, Any], *, deal=False):
    headers = [clean_text(x) for x in spec.get("headers", [])]
    rows = [[clean_text(x) for x in row] for row in spec.get("rows", [])]
    if deal:
        headers = []
        widths = spec.get("widths", [2100, PAGE_WIDTH_DXA - 2100])
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        set_table_geometry(table, widths)
        set_table_borders(table, color="808080", size="4", inside=False)
        for r_idx, row in enumerate(rows):
            for c_idx in range(2):
                cell = table.cell(r_idx, c_idx)
                shade_cell(cell, DEAL_LEFT if c_idx == 0 else DEAL_RIGHT)
                p = cell.paragraphs[0]
                set_para_spacing(p, after=0, line=1.02)
                add_runs_from_markdown(
                    p,
                    row[c_idx] if c_idx < len(row) else "",
                    size=8.8,
                    bold_default=(c_idx == 0),
                )
        keep_table_rows_together(table)
        doc.add_paragraph()
        return table

    widths = infer_widths(headers, rows, spec.get("widths"))
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table, color=BORDER, size="4")

    for c_idx, header in enumerate(headers):
        cell = table.cell(0, c_idx)
        shade_cell(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0, line=1.0)
        add_runs_from_markdown(p, header, size=8.9, bold_default=True)

    numeric_cols = set(spec.get("numeric_columns", []))
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            if c_idx in numeric_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_para_spacing(p, after=0, line=1.02)
            add_runs_from_markdown(p, value, size=8.65)
    keep_table_rows_together(table)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color="D6DCE4", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=0, line=1.08)
    lr = p.add_run(f"{clean_text(label)}: ")
    set_run_font(lr, size=10.2, bold=True)
    add_runs_from_markdown(p, clean_text(text), size=10.2)
    doc.add_paragraph()


def render_section(doc: Document, section: dict[str, Any]):
    if section.get("page_break_before"):
        doc.add_page_break()
    add_section_bar(doc, section.get("title", "Section"))

    for text in section.get("paragraphs", []):
        add_paragraph(doc, text)

    if "deal_table" in section:
        add_paragraph(doc, "The Deal:", bold=True, after=3)
        add_table(doc, section["deal_table"], deal=True)

    for item in section.get("bullets", []):
        add_bullet(doc, item)
    if section.get("bullets"):
        doc.add_paragraph()

    for table in section.get("tables", []):
        add_table(doc, table)

    for callout in section.get("callouts", []):
        add_callout(doc, callout.get("label", "Read"), callout.get("text", ""))


def build_memo(data: dict[str, Any], out_path: Path):
    doc = Document()
    configure_document(doc, data.get("footer", "Confidential memo"))
    add_masthead(doc, data)
    for section in data.get("sections", []):
        render_section(doc, section)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a house-style VC memo DOCX from JSON.")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    data = json.loads(args.input_json.read_text())
    build_memo(data, args.output_docx)
    print(f"Wrote {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
